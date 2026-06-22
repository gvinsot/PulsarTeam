"""Word (.docx) — read, edit-in-place, generate. Backed by python-docx.

Editing is round-trip: we load the existing document, mutate the parts the
operations target, and save — python-docx preserves everything we don't touch.
The LLM addresses paragraphs by integer index (see read_docx) and emits a list
of small JSON operations rather than rewriting the whole file.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from ..util import ToolError, check_size, default_output, resolve_output, resolve_path


def _para_md(p) -> str:
    style = (p.style.name if p.style else "") or ""
    text = p.text
    if style.startswith("Heading"):
        try:
            level = int(style.split()[-1])
        except ValueError:
            level = 1
        return f"{'#' * min(level, 6)} {text}"
    if style in ("List Bullet", "List Paragraph"):
        return f"- {text}"
    return text


def read_docx(path: str) -> dict[str, Any]:
    """Return paragraphs (indexed), tables, and a markdown rendering."""
    src = resolve_path(path)
    check_size(src)
    doc = Document(str(src))

    paragraphs = [
        {"index": i, "style": (p.style.name if p.style else None), "text": p.text}
        for i, p in enumerate(doc.paragraphs)
    ]
    tables = []
    for ti, t in enumerate(doc.tables):
        rows = [[c.text for c in row.cells] for row in t.rows]
        tables.append({"index": ti, "rows": rows})

    md_lines = [_para_md(p) for p in doc.paragraphs]
    return {
        "format": "docx",
        "path": str(src),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "markdown": "\n\n".join(line for line in md_lines if line is not None),
    }


def get_outline(path: str) -> dict[str, Any]:
    """Cheap navigation: just the heading hierarchy with paragraph indexes."""
    src = resolve_path(path)
    doc = Document(str(src))
    headings = []
    for i, p in enumerate(doc.paragraphs):
        name = p.style.name if p.style else ""
        if name and name.startswith("Heading") and p.text.strip():
            try:
                level = int(name.split()[-1])
            except ValueError:
                level = 1
            headings.append({"index": i, "level": level, "text": p.text})
    return {"format": "docx", "path": str(src), "headings": headings}


def _apply_op(doc, op: dict[str, Any]) -> str:
    kind = op.get("op")
    paras = doc.paragraphs
    if kind == "replace_text":
        find, repl = op.get("find"), op.get("replace", "")
        if not find:
            raise ToolError("replace_text requires 'find'")
        count = 0
        only_first = op.get("first", False)
        for p in paras:
            if find in p.text:
                # Rewrite via runs to keep paragraph-level formatting where possible.
                inline = p.runs
                if inline:
                    full = "".join(r.text for r in inline).replace(find, repl)
                    inline[0].text = full
                    for r in inline[1:]:
                        r.text = ""
                else:
                    p.text = p.text.replace(find, repl)
                count += 1
                if only_first:
                    break
        return f"replace_text: {count} paragraph(s)"
    if kind == "set_paragraph":
        idx = op["index"]
        if not 0 <= idx < len(paras):
            raise ToolError(f"set_paragraph index out of range: {idx}")
        paras[idx].text = op.get("text", "")
        return f"set_paragraph[{idx}]"
    if kind == "insert_paragraph":
        idx = op.get("after_index", len(paras) - 1)
        if not -1 <= idx < len(paras):
            raise ToolError(f"insert_paragraph after_index out of range: {idx}")
        anchor = paras[idx] if idx >= 0 else paras[0]
        new_p = anchor.insert_paragraph_before(op.get("text", "")) if idx < 0 else None
        if new_p is None:
            # Insert AFTER the anchor by inserting before its next sibling.
            from docx.oxml.ns import qn  # local import keeps top clean

            new_para = anchor._p.addnext(anchor._p.makeelement(qn("w:p"), {}))
            from docx.text.paragraph import Paragraph

            np = Paragraph(new_para, anchor._parent)
            np.text = op.get("text", "")
            if op.get("style"):
                np.style = op["style"]
            return f"insert_paragraph after {idx}"
        if op.get("style"):
            new_p.style = op["style"]
        return "insert_paragraph at start"
    if kind == "delete_paragraph":
        idx = op["index"]
        if not 0 <= idx < len(paras):
            raise ToolError(f"delete_paragraph index out of range: {idx}")
        p = paras[idx]
        p._p.getparent().remove(p._p)
        return f"delete_paragraph[{idx}]"
    raise ToolError(f"unknown docx op: {kind}")


def edit_docx(path: str, operations: list[dict[str, Any]], output_path: str | None = None) -> dict[str, Any]:
    """Apply a list of edit operations; save to output_path (default: *.edited.docx)."""
    src = resolve_path(path)
    doc = Document(str(src))
    if not isinstance(operations, list) or not operations:
        raise ToolError("operations must be a non-empty list")
    applied = [_apply_op(doc, op) for op in operations]
    out = resolve_output(output_path) if output_path else default_output(src)
    doc.save(str(out))
    return {"ok": True, "applied": applied, "output_path": str(out)}


def generate_docx(output_path: str, spec: dict[str, Any]) -> dict[str, Any]:
    """Create a new .docx from a structured spec.

    spec = {
      "title": "...",                         # optional H0/title
      "sections": [
        {"heading": "...", "level": 1,
         "paragraphs": ["...", ...],
         "bullets": ["...", ...],
         "table": {"header": [...], "rows": [[...], ...]}}
      ]
    }
    """
    out = resolve_output(output_path)
    doc = Document()
    if spec.get("title"):
        doc.add_heading(str(spec["title"]), level=0)
    for sec in spec.get("sections", []):
        if sec.get("heading"):
            doc.add_heading(str(sec["heading"]), level=int(sec.get("level", 1)))
        for para in sec.get("paragraphs", []) or []:
            doc.add_paragraph(str(para))
        for bullet in sec.get("bullets", []) or []:
            doc.add_paragraph(str(bullet), style="List Bullet")
        table = sec.get("table")
        if table and table.get("rows"):
            header = table.get("header")
            ncols = len(header) if header else len(table["rows"][0])
            t = doc.add_table(rows=0, cols=ncols)
            t.style = "Table Grid"
            if header:
                cells = t.add_row().cells
                for c, val in zip(cells, header):
                    c.text = str(val)
            for row in table["rows"]:
                cells = t.add_row().cells
                for c, val in zip(cells, row):
                    c.text = str(val)
    doc.save(str(out))
    return {"ok": True, "output_path": str(out)}
